"""
Script to update both DOCX documents with PyQUBO integration section.
Run from: E:\qubo\qubo final\final\docs\
"""
from docx import Document
from docx.shared import Pt


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def add_para(doc, text):
    return doc.add_paragraph(text)


def add_code(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return p


def append_pyqubo_section(doc):
    # -- Section 15 ------------------------------------------------------------
    add_heading(doc, "15.  PyQUBO Integration  -  The Enterprise Ecosystem Bridge", level=1)

    add_para(
        doc,
        "With Loucetius v2.5, the solver is no longer a standalone island. It integrates "
        "directly with PyQUBO  -  the industry-standard library for defining quantum and "
        "combinatorial optimisation problems  -  positioning Loucetius as a production-ready "
        "drop-in replacement for D-Wave, Toshiba SBM, and other commercial quantum platforms."
    )

    # 15.1
    add_heading(doc, "15.1  The Enterprise Adoption Problem", level=2)
    add_para(
        doc,
        "Enterprise quantum teams at aerospace firms, financial institutions, and logistics "
        "companies have invested years and thousands of lines of code into PyQUBO-based "
        "problem formulations. These teams write their QUBO objectives, constraints, and "
        "penalty terms in PyQUBO's high-level symbolic notation  -  and then submit the "
        "compiled model to a cloud quantum sampler."
    )
    add_para(
        doc,
        "Until now, switching to a new solver required rewriting every problem definition. "
        "This migration cost was a fundamental barrier to enterprise adoption. "
        "Loucetius removes this barrier entirely."
    )

    # 15.2
    add_heading(doc, "15.2  The Drop-In Replacement", level=2)
    add_para(
        doc,
        "The Loucetius PyQUBO API accepts any compiled PyQUBO model directly. "
        "A user who previously called D-Wave's LeapHybridSampler changes exactly one line:"
    )
    add_code(
        doc,
        "# BEFORE (D-Wave):\n"
        "answer = LeapHybridSampler().sample(pyqubo_model)\n\n"
        "# AFTER (Loucetius  -  one import change):\n"
        "from LOUCETIUS_pyqubo_api import sample\n"
        "answer = sample(pyqubo_model)"
    )
    add_para(
        doc,
        "The result is returned in a D-Wave-compatible SampleSet format. Downstream code "
        "that processes D-Wave results requires zero modification. The enterprise inherits "
        "Loucetius' full capability  -  200,000+ variable capacity, native GPU acceleration, "
        "Kerr and Hybrid engine routing  -  with no code rewrites whatsoever."
    )

    # 15.3
    add_heading(doc, "15.3  The Translation Layer (How It Works)", level=2)
    add_para(
        doc,
        "Internally, the PyQUBO API translates a compiled BinaryQuadraticModel (BQM) into "
        "a Loucetius sparse COO Q-matrix. The translation proceeds in three deterministic steps:"
    )
    add_para(
        doc,
        "Step 1  -  Variable Mapping: Each symbolic variable in the PyQUBO model is assigned "
        "a deterministic integer index. Variable names (strings such as 'asset_0', 'x_2_7') "
        "are sorted and mapped to matrix indices 0 through N-1. This determinism guarantees "
        "reproducible results across runs."
    )
    add_para(
        doc,
        "Step 2  -  Diagonal Population: Linear terms (single-variable coefficients h_i) are "
        "placed on the diagonal: Q[i, i] = h_i. These represent the individual energy "
        "contributions of each binary variable in isolation."
    )
    add_para(
        doc,
        "Step 3  -  Off-Diagonal Population: Quadratic interaction terms (two-variable "
        "couplings J_{ij}) are split symmetrically: Q[i, j] += J/2 and Q[j, i] += J/2. "
        "This symmetric form is required by the Loucetius C++ core and ensures correct "
        "energy evaluation under the standard QUBO formulation x^T Q x."
    )
    add_para(
        doc,
        "The resulting COO sparse matrix is passed directly to the Topographical Profiler, "
        "which routes it to the optimal physics engine as described in Section 14."
    )

    # 15.4
    add_heading(doc, "15.4  PyQUBO Advantages for Problem Definition", level=2)
    add_para(
        doc,
        "PyQUBO eliminates the laborious manual construction of QUBO matrices. Consider the "
        "difference: defining a Knapsack problem by hand requires manually computing slack "
        "variable encodings, penalty weights, and expanding symbolic quadratic expressions "
        "into explicit matrix entries. With PyQUBO, the user writes the constraint expression "
        "symbolically and the compiler handles all derivations automatically."
    )
    add_para(
        doc,
        "This separation of concerns is architecturally clean: PyQUBO handles the 'what' "
        "(problem definition) and Loucetius handles the 'how' (physical optimisation). "
        "Enterprise teams maintain their human-readable problem definitions while benefiting "
        "fully from Loucetius' GPU-native physics engine."
    )

    # 15.5
    add_heading(doc, "15.5  Engine Selection for PyQUBO Problems", level=2)
    add_para(
        doc,
        "When using the PyQUBO API, the engine parameter selects the physics backend. "
        "AUTO is recommended for most cases, but domain-specific guidance:"
    )
    add_para(
        doc,
        "KERR  -  Strictly constrained problems (financial portfolios with regulatory limits, "
        "aerospace routing with hard load constraints). The Kerr Wave Engine embeds constraint "
        "geometry directly into wave physics, architecturally preventing infeasible solutions."
    )
    add_para(
        doc,
        "HYBRID  -  Mixed problems combining broad landscape exploration with constraint "
        "enforcement (metamaterial design, logistics with soft and hard constraints). "
        "The Dynamic Momentum Blend adapts the Kerr-to-thermal ratio in real time based "
        "on the problem's constraint density profile."
    )
    add_para(
        doc,
        "SPARSE  -  Extremely large PyQUBO problems exceeding 10,000 variables (structural "
        "meshes, large graph problems, city-scale logistics). The cuSPARSE engine handles "
        "these using physical-connection mapping with full precision preservation."
    )

    # 15.6
    add_heading(doc, "15.6  Batch Processing and Parameter Sweeps", level=2)
    add_para(
        doc,
        "The sample_batch() function accepts a list of compiled PyQUBO models and solves them "
        "sequentially, returning a list of SampleSet objects. This is particularly useful for "
        "portfolio parameter sweeps, risk sensitivity analysis, or multi-scenario logistics "
        "optimisation  -  where the same structural problem must be evaluated under many different "
        "parameter conditions without rewriting the pipeline."
    )

    # 15.7
    add_heading(doc, "15.7  The Enterprise Value Proposition", level=2)
    add_para(
        doc,
        "The combined effect of the PyQUBO ecosystem bridge and Loucetius' physical "
        "architecture creates a compelling enterprise proposition:"
    )
    add_para(
        doc,
        "Capability Comparison: D-Wave Advantage 2000+ handles approximately 5,000 variables. "
        "Loucetius handles 200,000+ variables on standard RTX hardware, with no cloud "
        "dependency, no queue time, and no per-use cost."
    )
    add_para(
        doc,
        "Migration Cost: Zero. Existing PyQUBO problem definitions require no modification. "
        "One import statement changes. No code rewrites. No re-validation of problem formulations."
    )
    add_para(
        doc,
        "Physics Integrity: Unlike cloud quantum annealers with embedding overhead and "
        "hardware noise, Loucetius solves the exact problem as defined, with full precision "
        "and deterministic result formats. No embedding errors. No chain breaks. "
        "No topology-limited adjacency constraints."
    )


# -- MANUAL --------------------------------------------------------------------
man = Document("LOCETIUS_USER_MANUAL.docx")
append_pyqubo_section(man)
man.save("LOCETIUS_USER_MANUAL.docx")
print(f"Manual saved. Paragraphs: {len(man.paragraphs)}")

# -- USER GUIDE ----------------------------------------------------------------
guide = Document("Locetius_v2.5_User_Guide.docx")
append_pyqubo_section(guide)
guide.save("Locetius_v2.5_User_Guide.docx")
print(f"User Guide saved. Paragraphs: {len(guide.paragraphs)}")
